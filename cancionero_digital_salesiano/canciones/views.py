from django.contrib import messages
from django.shortcuts import redirect, render, get_object_or_404
from datetime import date, timedelta
from .models import Cancion, Favorito, ListaCancion, ListaPersonal, TiempoLiturgico, ListaCancion
from django.template.loader import render_to_string
from django.http import JsonResponse, HttpResponse
import re
from django.contrib.auth.decorators import login_required
from docx import Document

# ============================
# 🗓 CÁLCULOS DE FECHAS Y TIEMPOS LITÚRGICOS
# ============================


def calcular_fecha_pascua(anio):
    """
    Calcula la fecha de Pascua para un año dado según el algoritmo Computus.
    Retorna un objeto date con la fecha exacta de Pascua para ese año.
    """
    a = anio % 19
    b = anio // 100
    c = anio % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    mes = (h + l - 7 * m + 114) // 31
    dia = ((h + l - 7 * m + 114) % 31) + 1
    return date(anio, mes, dia)


def obtener_tiempo_liturgico_actual():
    """
    Devuelve el nombre del tiempo litúrgico actual basado en la fecha actual.
    Calcula los rangos de fechas para Adviento, Navidad, Cuaresma, Pascua, y Tiempo Ordinario.
    Retorna el nombre del tiempo litúrgico activo o 'Ampliación' si no coincide con ninguno.
    """
    hoy = date.today()
    anio = hoy.year

    # Calculamos fechas clave usando la fecha de Pascua
    pascua = calcular_fecha_pascua(anio)
    cuaresma_inicio = pascua - timedelta(days=46)
    cuaresma_fin = pascua - timedelta(days=1)
    pascua_fin = pascua + timedelta(days=49)  # Hasta Pentecostés

    navidad_inicio = date(anio, 12, 25)
    bautismo_senor = navidad_inicio + timedelta(days=19)
    tiempo_ordinario_1 = bautismo_senor + timedelta(days=1)
    tiempo_ordinario_2 = pascua_fin + timedelta(days=1)

    # Calculo del cuarto domingo antes de Navidad (inicio de Adviento)
    cuarto_domingo_antes_navidad = navidad_inicio - timedelta(
        days=((navidad_inicio.weekday() + 1) % 7) + 21
    )
    adviento_inicio = cuarto_domingo_antes_navidad
    cristo_rey = adviento_inicio - timedelta(days=7)

    # Determinamos en qué periodo litúrgico estamos
    if adviento_inicio <= hoy < navidad_inicio:
        return "Adviento"
    elif navidad_inicio <= hoy <= bautismo_senor:
        return "Navidad"
    elif cuaresma_inicio <= hoy <= cuaresma_fin:
        return "Cuaresma"
    elif pascua <= hoy <= pascua_fin:
        return "Pascua"
    elif tiempo_ordinario_1 <= hoy < cuaresma_inicio:
        return "Tiempo Ordinario"
    elif tiempo_ordinario_2 <= hoy < adviento_inicio:
        return "Tiempo Ordinario"
    else:
        return "Ampliación"  # Valor por defecto si no coincide con los anteriores


def obtener_tiempo_liturgico(id_tiempo):
    """
    Recupera el objeto TiempoLiturgico por su id.
    Si no existe, lanza error 404.
    """
    return get_object_or_404(TiempoLiturgico, id_tiempo=id_tiempo)


# ============================
# 🏠 VISTAS PRINCIPALES DE LA APLICACIÓN
# ============================


def index(request):
    """
    Vista principal que muestra las canciones correspondientes al tiempo litúrgico actual.
    Si no se encuentra, muestra las canciones del tiempo 'Ampliación' por defecto.
    """
    nombre_tiempo = obtener_tiempo_liturgico_actual()
    tiempo_actual = TiempoLiturgico.objects.filter(
        nombre_tiempo__iexact=nombre_tiempo
    ).first()

    if not tiempo_actual:
        # Si no se encuentra el tiempo litúrgico calculado, usa 'Ampliación'
        tiempo_actual = TiempoLiturgico.objects.filter(
            nombre_tiempo__iexact="Ampliación"
        ).first()

    # Obtener canciones filtradas por el tiempo litúrgico actual
    canciones = (
        Cancion.objects.filter(id_tiempo=tiempo_actual.id_tiempo)
        if tiempo_actual
        else Cancion.objects.none()
    )

    return render(
        request,
        "canciones/index.html",
        {
            "tiempo_actual": tiempo_actual,
            "canciones": canciones,
        },
    )


def canciones_complete(request):
    tiempos = TiempoLiturgico.objects.all()  # Todos los tiempos litúrgicos disponibles

    canciones_por_tiempo = {}  # Diccionario para almacenar canciones agrupadas

    for tiempo in tiempos:
        canciones_por_tiempo[tiempo.nombre_tiempo] = Cancion.objects.filter(
            id_tiempo=tiempo.id_tiempo
        )

    return render(
        request,
        "canciones/toggleCanciones.html",
        {
            "canciones_por_tiempo": canciones_por_tiempo,
            "tiempos": tiempos,
        },
    )


"""
@login_required
def lista_favoritos(request):
    # Obtener las canciones favoritas del usuario
    favoritos = Favorito.objects.filter(usuario=request.user).select_related('cancion')
    canciones_favoritas = [f.favorito for f in favoritos]  # Obtener solo las canciones, no los objetos Favorito completos
    return render(request, 'favoritos/lista_favoritos.html', {'canciones': canciones_favoritas})
"""
# ============================
# 🎸 FUNCIONES DE TRANSPOSE DE ACORDES
# ============================

# Lista de notas musicales en español con sostenidos (#)
NOTAS = [
    "DO", "DO#", "RE", "RE#", "MI", "FA", "FA#", "SOL", "SOL#", "LA", "LA#", "SI"
]

# Diccionario para notas con dobles sostenidos o notas raras que se deben mapear a enharmónicas comunes
ENHARMONICAS_SIMPLIFICADAS = {
    "DO##": "RE",
    "RE##": "MI",
    "FA##": "SOL",
    "SOL##": "LA",
    "LA##": "SI",
    "MI#": "FA",
    "SI#": "DO",
    "DOB": "DO#",
    "REB": "RE#",
    "MIB": "MI",
    "SOLB": "FA#",
    "LAB": "SOL#",
    "SIB": "LA#"
}


# Diccionario para notas bemoles equivalentes (enharmónicas)
BEMOLES = {
    "DOB": "DO#",
    "REB": "RE#",
    "MIB": "MI",
    "SOLB": "FA#",
    "LAB": "SOL#",
    "SIB": "LA#",
    "B": "SI"  # caso raro pero para ilustrar
}

# Función para obtener índice en lista NOTAS
def obtener_indice_nota(nota):
    nota_upper = nota.upper()
    if nota_upper in NOTAS:
        return NOTAS.index(nota_upper)
    elif nota_upper in BEMOLES:
        return NOTAS.index(BEMOLES[nota_upper])
    else:
        return -1

def mantener_case(original, nuevo):
    # Si el original es todo mayúsculas, devuelve todo mayúsculas
    if original.isupper():
        return nuevo.upper()
    # Si la primera letra es mayúscula, el resto minúsculas
    elif original[0].isupper():
        return nuevo.capitalize()
    # Sino minúsculas
    else:
        return nuevo.lower()

def transponer_acorde(acorde, semitonos):
    import re

    pattern = r'^([a-zA-Z]+[#b]?)(.*)$'
    match = re.match(pattern, acorde)
    if not match:
        return acorde

    base = match.group(1)
    sufijo = match.group(2)

    idx = obtener_indice_nota(base)
    if idx == -1:
        return acorde

    nuevo_idx = (idx + semitonos) % len(NOTAS)
    nueva_base = NOTAS[nuevo_idx]

    # Simplificar enharmónicas raras si aparecen
    # Ej: fa## => sol
    if nueva_base.upper() in ENHARMONICAS_SIMPLIFICADAS:
        nueva_base = ENHARMONICAS_SIMPLIFICADAS[nueva_base.upper()]

    nueva_base = mantener_case(base, nueva_base)

    return nueva_base + sufijo


def transponer_linea(linea, semitonos=1):
    # Regex para encontrar acordes en la línea (simplificado para tu contexto)
    pattern = r'\b([a-zA-Z]+[#b]?m?7?|maj7|sus4|sus2|6|m6)\b'

    def reemplazo(match):
        acorde = match.group(0)
        return transponer_acorde(acorde, semitonos)

    return re.sub(pattern, reemplazo, linea)

# ============================
# 🎵 VISTA DETALLE DE CANCIÓN CON TRANSPOSE
# ============================

def detalle_cancion(request, id_cancion):
    cancion = get_object_or_404(Cancion, id_cancion=id_cancion)
    semitonos = int(request.GET.get("transpose", 0))

    lineas = cancion.lineacancion_set.all().order_by("linea_num")
    lineas_transpuestas = []
    for linea in lineas:
        contenido = linea.contenido
        if linea.tipo_linea == "acorde" and semitonos != 0:
            contenido = transponer_linea(contenido, semitonos)
        lineas_transpuestas.append(
            {"contenido": contenido, "tipo_linea": linea.tipo_linea}
        )

    # Fragmento dinámico para AJAX
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        html = render_to_string(
            "canciones/lineas_fragment.html", {"lineas": lineas_transpuestas}
        )
        return JsonResponse({"html": html})

    # Datos adicionales
    if request.user.is_authenticated:
        favorito = Favorito.objects.filter(usuario=request.user, cancion=id_cancion)
    else:
        favorito = None  # o un queryset vacío: Favorito.objects.none()

    tiempo_actual = TiempoLiturgico.objects.get(id_tiempo=cancion.id_tiempo.id_tiempo)  # Lógica de tiempo litúrgico
    listas_usuario = []
    listas_asociadas = []
    if request.user.is_authenticated:
        listas_usuario = ListaPersonal.objects.filter(usuario=request.user)
        listas_asociadas = ListaPersonal.objects.filter(
        listacancion__cancion=cancion, usuario=request.user
    ).distinct()

    return render(
        request,
        "canciones/cancion.html",
        {
            "cancion": cancion,
            "lineas": lineas_transpuestas,
            "semitonos": semitonos,
            "favorito": favorito,
            "tiempo_actual": tiempo_actual,
            "listas_usuario": listas_usuario,
            "listas_asociadas": listas_asociadas,
        },
    )

# ============================
# 🔍 FUNCIONALIDAD DE BÚSQUEDA
# ============================


def search(request):
    """
    Busca canciones por título que contengan la cadena 'q' enviada por GET.
    Solo realiza búsqueda si la longitud de la consulta es >= 3.
    Si es petición AJAX, devuelve fragmento HTML con resultados.
    """
    query = request.GET.get("q", "")
    nombre_tiempo = obtener_tiempo_liturgico_actual()
    tiempo_actual = TiempoLiturgico.objects.filter(
        nombre_tiempo__iexact=nombre_tiempo
    ).first()

    if len(query) >= 3:
        canciones = Cancion.objects.filter(titulo__icontains=query)
    else:
        # Si la búsqueda es muy corta, muestra canciones del tiempo actual
        canciones = (
            Cancion.objects.filter(id_tiempo=tiempo_actual.id_tiempo)
            if tiempo_actual
            else Cancion.objects.none()
        )

    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        html = render_to_string(
            "canciones/canciones_list.html", {"canciones": canciones}
        )
        return JsonResponse({"html": html})

    return render(
        request,
        "canciones/index.html",
        {
            "tiempo_actual": tiempo_actual,
            "canciones": canciones,
            "busqueda": query,
        },
    )


@login_required
def profile(request):
    return render(request, "account/profile.html")


# POLITICAS
def politica_privacidad(request):
    return render(
        request,
        "canciones/info_pagina.html",
        {
            "titulo": "Política de Privacidad",
            "contenido": """
        En Cancionero Digital Salesiano, respetamos tu privacidad y nos comprometemos a proteger los datos personales que nos proporciones.

        ### 1. Recopilación de información
        Recopilamos datos como nombre de usuario y dirección de correo electrónico cuando te registras en la plataforma. Esta información se utiliza únicamente con fines de autenticación y personalización de tu experiencia.

        ### 2. Uso de la información
        Utilizamos los datos recopilados para ofrecerte acceso a funcionalidades como listas personalizadas, favoritos y ajustes de usuario. No compartimos esta información con terceros.

        ### 3. Seguridad de los datos
        Implementamos medidas de seguridad técnicas y organizativas para proteger tu información frente a accesos no autorizados.

        ### 4. Derechos del usuario
        Puedes solicitar el acceso, rectificación o eliminación de tus datos escribiéndonos a: contacto@cancionerosalesiano.com

        Esta política es ficticia y se proporciona únicamente con fines académicos.
        """,
        },
    )


def aviso_legal(request):
    return render(
        request,
        "canciones/info_pagina.html",
        {
            "titulo": "Aviso Legal",
            "contenido": """
        Este sitio web, "Cancionero Digital Salesiano", es un proyecto académico desarrollado como parte de un Trabajo de Fin de Grado.

        ### 1. Titularidad del sitio
        El sitio está gestionado por un estudiante de Desarrollo de Aplicaciones Web con fines educativos. No representa a ninguna entidad oficial ni comercial.

        ### 2. Propiedad intelectual
        Todos los contenidos, incluyendo canciones, textos, imágenes o código, son utilizados únicamente para demostración. Si algún contenido infringe derechos, será eliminado inmediatamente previa notificación.

        ### 3. Responsabilidad
        No nos hacemos responsables del uso indebido que se pueda hacer de la información contenida en este sitio.

        ### 4. Contacto
        Para cualquier reclamación o consulta, puedes escribir a: contacto@cancionerosalesiano.com

        Este aviso es simulado para propósitos académicos.
        """,
        },
    )


def politica_cookies(request):
    return render(
        request,
        "canciones/info_pagina.html",
        {
            "titulo": "Política de Cookies",
            "contenido": """
        Esta web utiliza cookies propias y de terceros para mejorar la experiencia del usuario. Dado que es un proyecto académico, el uso de cookies es mínimo y no se utilizan con fines comerciales.

        ### 1. ¿Qué son las cookies?
        Las cookies son pequeños archivos que se almacenan en tu navegador cuando visitas una web.

        ### 2. Tipos de cookies utilizadas
        - Cookies técnicas: Necesarias para el funcionamiento del sitio.
        - Cookies de sesión: Para mantener la sesión iniciada mientras navegas.

        ### 3. Gestión de cookies
        Puedes eliminar o bloquear las cookies desde tu navegador. Ten en cuenta que esto podría afectar al funcionamiento del sitio.

        Este texto es meramente demostrativo para un trabajo académico.
        """,
        },
    )


def terminos_condiciones(request):
    return render(
        request,
        "canciones/info_pagina.html",
        {
            "titulo": "Términos y Condiciones",
            "contenido": """
        Bienvenido al sitio web "Cancionero Digital Salesiano". Al acceder y utilizar esta plataforma, aceptas los siguientes términos:

        ### 1. Uso permitido
        Este sitio está diseñado exclusivamente para uso académico y personal. Está prohibido su uso con fines comerciales o ilícitos.

        ### 2. Registro y cuentas
        Los usuarios registrados son responsables de la veracidad de los datos proporcionados y del uso adecuado de su cuenta.

        ### 3. Propiedad del contenido
        Las canciones y listas presentadas son ejemplos educativos. El desarrollador no reclama derechos sobre ellas.

        ### 4. Modificaciones
        Nos reservamos el derecho de modificar estos términos sin previo aviso. Las versiones actualizadas se publicarán en esta misma página.

        Este texto forma parte de una simulación legal para un Trabajo de Fin de Grado.
        """,
        },
    )


@login_required
def toggle_favorito(request):
    # Asegurarse de que la solicitud sea tipo POST
    if request.method == "POST":
        # Obtener el ID de la canción desde el formulario
        cancion_id = request.POST.get("cancion_id")
        # Obtener la canción desde la base de datos, o lanzar 404 si no existe
        cancion = get_object_or_404(Cancion, id_cancion=cancion_id)

        # Verificar si ya existe un favorito del usuario para esta canción
        favorito = Favorito.objects.filter(
            usuario=request.user, cancion=cancion
        ).first()

        if favorito:
            # Si ya existe, eliminarlo (quitar de favoritos)
            favorito.delete()
            return JsonResponse({"status": "eliminado"})
        else:
            # Si no existe, agregarlo como favorito
            Favorito.objects.create(usuario=request.user, cancion=cancion)
            return JsonResponse({"status": "agregado"})

    # Si la solicitud no es POST, devolver error 405 (método no permitido)
    return JsonResponse({"error": "Método no permitido"}, status=405)

@login_required
def favoritos(request):
    # Obtener todos los registros de canciones favoritas del usuario
    favoritos_usuario = Favorito.objects.filter(usuario=request.user).select_related("cancion")

    # Extraer solo las instancias de canciones desde los favoritos
    canciones = [fav.cancion for fav in favoritos_usuario]

    # Renderizar plantilla con la lista de canciones favoritas
    return render(
        request,
        "canciones/favoritos.html",
        {
            "canciones": canciones,
        },
    )

@login_required
def toggle_list(request):
    # Verifica que sea una solicitud POST
    if request.method == "POST":
        # Obtener datos del formulario
        cancion_id = request.POST.get("cancion_id")
        lista_id = request.POST.get("lista_id")
        nueva_lista_nombre = request.POST.get("nueva_lista")

        # Obtener la canción correspondiente
        cancion = get_object_or_404(Cancion, id_cancion=cancion_id)

        # Si se indicó una nueva lista, crearla (o usar la existente)
        if nueva_lista_nombre:
            lista, created = ListaPersonal.objects.get_or_create(
                nombre=nueva_lista_nombre, usuario=request.user
            )
        else:
            # Si no, obtener la lista existente con el ID dado
            lista = get_object_or_404(ListaPersonal, id=lista_id, usuario=request.user)

        # Agregar o quitar la canción de la lista según si ya está o no
        if cancion in lista.canciones.all():
            lista.canciones.remove(cancion)
            return JsonResponse({"status": "eliminado"})
        else:
            lista.canciones.add(cancion)
            return JsonResponse({"status": "agregado"})

    # Si la solicitud no es POST, devolver error
    return JsonResponse({"error": "Método no permitido"}, status=405)


@login_required
def guardar_cancion_en_lista(request, cancion_id):
    # Solo acepta solicitudes POST
    if request.method == "POST":
        # Recoger datos del formulario
        lista_id = request.POST.get("lista_id")
        nueva_lista_nombre = request.POST.get("nueva_lista")

        # Crear o buscar la lista personal según si se ingresó un nuevo nombre
        if nueva_lista_nombre:
            lista, created = ListaPersonal.objects.get_or_create(
                usuario=request.user, nombre_lista=nueva_lista_nombre
            )
        else:
            lista = get_object_or_404(
                ListaPersonal, id_lista=lista_id, usuario=request.user
            )

        # Obtener la canción y asociarla a la lista (si no existe ya)
        cancion = get_object_or_404(Cancion, id_cancion=cancion_id)
        ListaCancion.objects.get_or_create(lista=lista, cancion=cancion)

        # Redirigir al detalle de la canción
        return redirect("detalle_cancion", id_cancion=cancion_id)

def lista(request):
    """
    Vista que muestra todas las listas personales del usuario con sus canciones.
    """
    
    if request.user.is_authenticated:
        # Obtener todas las listas creadas por el usuario autenticado
        listas = ListaPersonal.objects.filter(usuario=request.user)

        # Construir una lista de diccionarios con cada lista y sus canciones asociadas
        listas_con_canciones = []
        for lista in listas:
            # Obtener las canciones relacionadas con esta lista desde la tabla intermedia ListaCancion
            canciones = Cancion.objects.filter(listacancion__lista=lista)

            # Agregar al resultado
            listas_con_canciones.append(
                {
                    "lista": lista,
                    "canciones": canciones,
                }
            )


        # Renderizar la plantilla pasando la estructura de listas con canciones
        return render(
            request,
            "canciones/lista.html",
            {
                "listas_con_canciones": listas_con_canciones,
                "usuario": request.user,
                "listas": listas,
            },
        )
    else:
        # Si el usuario no está autenticado, redirigir a la página de inicio de sesión
        messages.error(request, "Debes iniciar sesión para ver tus listas.")
        return redirect("account_login")  

@login_required
def lista_detalle(request, id_lista):
    """
    Vista que muestra los detalles de una lista específica, incluyendo las canciones asociadas.
    """

    # Obtener la lista que pertenece al usuario o lanzar 404 si no existe o no es suya
    lista = get_object_or_404(ListaPersonal, id_lista=id_lista, usuario=request.user)

    # Obtener las canciones asociadas a esta lista
    canciones = Cancion.objects.filter(listacancion__lista=lista)

    # Renderizar la vista detalle de la lista con sus canciones
    return render(
        request,
        "canciones/lista_detalle.html",
        {
            "lista": lista,
            "canciones": canciones,
        },
    )

@login_required
def exportar_lista_word(request, id_lista):
    # Obtener la lista específica que pertenece al usuario autenticado
    lista = get_object_or_404(ListaPersonal, id_lista=id_lista, usuario=request.user)

    # Comprobar si se deben incluir acordes (marcado como parámetro en la URL)
    incluir_acordes = request.GET.get('acordes') == '1'

    # Obtener todas las canciones de la lista (usando la relación intermedia)
    canciones = lista.listacancion_set.select_related('cancion').all()

    # Crear el documento Word
    doc = Document()
    doc.add_heading(f'Lista: {lista.nombre_lista}', level=1)

    if canciones:
        for lc in canciones:
            cancion = lc.cancion
            doc.add_heading(cancion.titulo, level=2)

            # Obtener las líneas de la canción ordenadas por número
            lineas = cancion.lineacancion_set.order_by('linea_num')

            for linea in lineas:
                # Saltar líneas de acordes si el usuario eligió no incluirlos
                if linea.tipo_linea == 'acorde' and not incluir_acordes:
                    continue

                # Agregar línea al documento
                p = doc.add_paragraph()
                run = p.add_run(linea.contenido)

                # Aplicar formato (negrita) a acordes y estribillos
                if linea.tipo_linea in ['acorde', 'estribillo']:
                    run.bold = True

            # Agregar espacio entre canciones
            doc.add_paragraph('')
    else:
        # Si la lista está vacía, indicarlo en el documento
        doc.add_paragraph('No hay canciones en esta lista.')

    # Preparar respuesta HTTP para descargar el documento
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"{lista.nombre_lista.replace(' ', '_')}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'

    # Guardar el documento en la respuesta y devolverlo
    doc.save(response)
    return response


#Eliminar una lista personal
def eliminar_lista(request, id_lista):
    lista = get_object_or_404(ListaPersonal, id_lista=id_lista, usuario=request.user)
    if request.method == 'POST':
        lista.delete()
        messages.success(request, 'La lista ha sido eliminada con éxito.')
        return redirect('lista')  # Ajusta el nombre a tu vista de listas
    return redirect('detalle_lista', id_lista=id_lista)
